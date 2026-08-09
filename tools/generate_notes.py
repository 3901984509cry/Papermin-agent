# =============================================================================
# tools/generate_notes.py — 工具④：自动生成读书笔记 Word 文档
# 功能：将 Agent 整理的内容写入格式化的 .docx 文件（真实文件输出）
# =============================================================================

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain.tools import Tool
from config import OUTPUT_DIR


def _set_heading_style(paragraph, text: str, level: int = 1):
    """设置标题样式"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    run.bold      = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = RGBColor(0x0A, 0x25, 0x40) if level == 1 else RGBColor(0x0D, 0x94, 0x88)


def generate_notes_doc(content: str) -> str:
    """
    根据 Agent 提供的内容生成结构化 Word 读书笔记。

    Args:
        content: 笔记内容，使用 ## 分隔章节，格式示例：
                 ## 研究背景
                 这篇论文的背景是...
                 ## 核心方法
                 作者提出了...
    Returns:
        生成结果描述（含文件路径）
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # ── 文档基础设置 ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"   # ⚠️ 如果系统无此字体可改为 "SimSun" 或 "Arial"
    style.font.size = Pt(11)

    # ── 封面标题 ──────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("📘 PaperMind 智能读书笔记")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x0A, 0x25, 0x40)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")

    doc.add_paragraph()  # 空行

    # ── 解析并写入内容 ────────────────────────────────────────────────────────
    sections = content.split("##")
    for section in sections:
        section = section.strip()
        if not section:
            continue

        lines = section.split("\n")
        section_title = lines[0].strip()
        section_body  = "\n".join(lines[1:]).strip()

        # 写章节标题
        heading = doc.add_paragraph()
        heading.add_run(f"◆ {section_title}")
        _set_heading_style(heading, section_title, level=1)
        doc.add_paragraph()  # 标题后空行

        # 写正文（按段落写入）
        for para_text in section_body.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            # 识别要点（以 - 或 • 开头）
            if para_text.startswith(("-", "•", "*")):
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(para_text.lstrip("-•* "))
            else:
                para = doc.add_paragraph(para_text)
                para.paragraph_format.first_line_indent = Pt(22)  # 首行缩进

        doc.add_paragraph()  # 章节间空行

    # ── 页脚：免责声明 ─────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_run  = footer_para.add_run("⚠️ 本笔记由 PaperMind AI 自动生成，关键数据请以论文原文为准。")
    footer_run.font.size  = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 保存文件 ──────────────────────────────────────────────────────────────
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(OUTPUT_DIR, f"读书笔记_{timestamp}.docx")
    doc.save(output_path)

    return (
        f"✅ 读书笔记已生成！\n"
        f"   📁 文件路径：{os.path.abspath(output_path)}\n"
        f"   📖 包含章节：{len([s for s in content.split('##') if s.strip()])} 个\n"
        f"   💡 提示：用 Word 或 WPS 打开文件查看完整内容"
    )


generate_notes_tool = Tool(
    name="generate_notes_doc",
    func=generate_notes_doc,
    description=(
        "将整理好的读书笔记内容生成为 Word（.docx）文件并保存到本地。"
        "输入：结构化笔记内容，用 ## 分隔各章节，例如：\n"
        "'## 研究背景\n内容...\n## 核心方法\n内容...'\n"
        "用户说'生成笔记'、'保存为文档'、'导出笔记'时调用此工具。"
    ),
)
